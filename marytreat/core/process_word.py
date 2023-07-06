import os
from glob import iglob
from shutil import copy2

from docx import Document

from marytreat.core import local as l


def clear_drawing_descriptions(document):
    for paragraph in document.paragraphs:
        style = paragraph.style
        if style.name == 'Graphic':
            drawing = paragraph.runs[0].element[1]  # <w:drawing>
            for tag in drawing.iter():
                description = tag.attrib.get('descr')
                if description:
                    tag.set('descr', '')
    return document


def show_table_entries(document):
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                l.logger.info(cell)
                for paragraph in cell.paragraphs:
                    l.logger.info(paragraph.text)


def backup_doc(path):
    backup_doc_path = '.'.join([os.path.splitext(path)[0], '.bak'])
    copy2(path, backup_doc_path)


def get_xml_content(document):
    root = document.element[0]
    return l.XMLContent(root)


def prepare_for_batch_converter(path):
    backup_doc(path)
    document = Document(path)
    clear_drawing_descriptions(document)


def get_ditamap(proj_folder):
    for file in iglob(proj_folder + '/*.ditamap'):
        if file.endswith('.ditamap'):
            word_ditamap = l.LocalMap(file)
            word_ditamap.image_folder = l.os.path.join(proj_folder, 'media')
            return word_ditamap


def after_conversion(project_folder):
    ditamap = get_ditamap(project_folder)
    ditamap.create_root_concept()

    docdetails = ditamap.topics[0]  # assume it's the first topic in the map
    docdetails.format_docdetails_from_word()

    for t in ditamap.topics:
        t.content.wrap_images_in_fig()
        t.content.process_notes()
        t.content.create_shortdesc_from_first_p()
        t.write()

    for t in ditamap.topics:
        t.content.images_to_png()
        t.write()

    ditamap.add_topic_groups()
    ditamap.write()


if __name__ == "__main__":
    doc_folder = r'C:\hp_cheetahr5\TS5ES-00009 - Feeder Service'
    # doc_path = os.path.join(doc_folder, 'TS5ES-00017.docx')
    # prepare_for_batch_converter(doc_path)

    dita_project_folder = l.os.path.join(doc_folder, 'output - Copy')
    after_conversion(dita_project_folder)

    # ditamap = get_ditamap(dita_project_folder)
    # topic = ditamap.topics[10]
    # topic.content.create_shortdesc_from_first_p()
    # for t in ditamap.topics:
    #     t.content.create_shortdesc_from_first_p()
    #     break

# Convert with Oxygen Batch Converter using Oxygen 25.
# Create root concept.
# Process docdetails.
# Put images in 'fig's and get them out of single-cell tables.
# Wrap steps.
# Update topic groups in map.
# Process notes.
# First paragraph becomes short description.
# Copy folder.
# Run MaryTreat, and it will recognize the types and update the map.

# 7. Add map metadata, "outputclass=body", "outputclass=appendixes" etc...
#
# Think about it:
# rewrite marytreat ui so that it does the same to docx-derived maps
# generate ISH files so that topic/map types don't get lost on the way to the server
# Desired result: feed this to MaryTreat and be happy, then feed this to Content Manager
